"""
claude-chat-app/utils/file_handler.py

In-memory file processor that converts user-uploaded files into Claude API
content blocks.

What it does:
  - Accepts any Streamlit UploadedFile object and returns a dict formatted for
    the OpenAI-compatible messages API (either an image_url block for images or
    a text block for all other types).
  - Detects file type by MIME type and extension, then dispatches to the
    appropriate extractor:
      Images  → base64-encoded image_url block for Claude Vision.
      PDF     → text extracted page-by-page with page markers (via PyPDF2).
      DOCX    → paragraphs and tables converted to Markdown (via python-docx).
      XLSX    → all sheets converted to Markdown tables (via openpyxl).
      CSV/TSV → first 200 rows as a Markdown table, with truncation notice.
      Code    → raw text wrapped in a fenced code block with language tag.
      Config  → raw text in a fenced block; JSON is pretty-printed first.
      Text    → raw UTF-8 text.
  - Enforces a 25 MB per-file limit (20 MB for images) and a 100,000-character
    context limit per file (truncates with a notice if exceeded).
  - No files are written to disk — all processing happens in memory within the
    request cycle. Files do not persist between Streamlit reruns.

Main entry point:
  prepare_file_for_api(uploaded_file) → dict  — call once per uploaded file;
  append the returned dict to the api_content list before the text prompt.
"""
import base64
import io
import os
import json

from PyPDF2 import PdfReader

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


# ===== SUPPORTED FILE TYPE MAPPINGS =====

CODE_EXTENSIONS = {
    ".py": "python", ".js": "javascript", ".ts": "typescript",
    ".java": "java", ".cpp": "cpp", ".c": "c", ".cs": "csharp",
    ".go": "go", ".rs": "rust", ".rb": "ruby", ".php": "php",
    ".swift": "swift", ".kt": "kotlin", ".scala": "scala",
    ".sh": "bash", ".bash": "bash", ".zsh": "zsh",
    ".html": "html", ".css": "css", ".scss": "scss",
    ".jsx": "jsx", ".tsx": "tsx", ".vue": "vue",
    ".sql": "sql", ".r": "r", ".m": "matlab",
    ".pl": "perl", ".lua": "lua", ".dart": "dart",
    ".tf": "terraform", ".hcl": "hcl",
}

CONFIG_EXTENSIONS = {
    ".json": "json", ".yaml": "yaml", ".yml": "yaml",
    ".toml": "toml", ".ini": "ini", ".cfg": "config",
    ".xml": "xml", ".env": "env", ".properties": "properties",
    ".ovpn": "openvpn", ".conf": "config",
    ".dockerfile": "dockerfile", ".gitignore": "gitignore",
    ".dockerignore": "dockerignore",
}

TEXT_EXTENSIONS = {
    ".txt": "text", ".md": "markdown", ".rst": "restructuredtext",
    ".log": "log", ".tex": "latex", ".rtf": "richtext",
}

MAX_FILE_SIZE_MB = 25
MAX_IMAGE_SIZE_MB = 20
MAX_CONTEXT_CHARS = 100000


# ===== TEXT EXTRACTION FUNCTIONS =====

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF with page markers."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- Page {i} ---\n{text}")
        return "\n\n".join(pages)
    except Exception as e:
        return f"[Error reading PDF: {str(e)}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word documents including tables."""
    if Document is None:
        return "[python-docx not installed - cannot read .docx files]"
    try:
        doc = Document(io.BytesIO(file_bytes))
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        content.append(f'{"#" * int(level)} {para.text}')
                    except ValueError:
                        content.append(f"## {para.text}")
                else:
                    content.append(para.text)
        for table in doc.tables:
            table_text = []
            for trow in table.rows:
                row_text = " | ".join(c.text.strip() for c in trow.cells)
                table_text.append(row_text)
            if table_text:
                content.append(table_text[0])
                content.append(" | ".join("---" for _ in table.rows[0].cells))
                content.extend(table_text[1:])
        return "\n".join(content) if content else "[Empty document]"
    except Exception as e:
        return f"[Error reading DOCX: {str(e)}]"


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract text from Excel files, all sheets."""
    if openpyxl is None:
        return "[openpyxl not installed - cannot read .xlsx files]"
    try:
        wb = openpyxl.load_workbook(
            io.BytesIO(file_bytes), read_only=True, data_only=True
        )
        all_sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in row_vals):
                    rows.append(" | ".join(row_vals))
            if rows:
                sheet_text = f"## Sheet: {sheet_name}\n"
                sheet_text += rows[0] + "\n"
                sheet_text += " | ".join("---" for _ in rows[0].split(" | ")) + "\n"
                sheet_text += "\n".join(rows[1:])
                all_sheets.append(sheet_text)
        wb.close()
        return "\n\n".join(all_sheets) if all_sheets else "[Empty spreadsheet]"
    except Exception as e:
        return f"[Error reading XLSX: {str(e)}]"


def extract_text_from_csv(file_bytes: bytes) -> str:
    """Extract text from CSV/TSV files as markdown table."""
    try:
        text = file_bytes.decode("utf-8")
        lines = text.strip().split("\n")
        if not lines:
            return "[Empty CSV file]"
        first_line = lines[0]
        delimiter = "\t" if "\t" in first_line else ","
        result = []
        for i, line in enumerate(lines[:200]):
            cells = line.split(delimiter)
            result.append(" | ".join(cells))
            if i == 0:
                result.append(" | ".join("---" for _ in cells))
        if len(lines) > 200:
            result.append(f"\n... ({len(lines) - 200} more rows truncated)")
        return "\n".join(result)
    except Exception as e:
        return f"[Error reading CSV: {str(e)}]"


def extract_text_from_code(file_bytes: bytes, filename: str) -> str:
    """Extract code files with syntax language tag."""
    try:
        text = file_bytes.decode("utf-8")
        ext = os.path.splitext(filename)[1].lower()
        lang = CODE_EXTENSIONS.get(ext, "")
        return f"```{lang}\n{text}\n```"
    except UnicodeDecodeError:
        return "[Binary file - cannot display as text]"


def extract_text_from_config(file_bytes: bytes, filename: str) -> str:
    """Extract config files, pretty-print JSON."""
    try:
        text = file_bytes.decode("utf-8")
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".json":
            try:
                parsed = json.loads(text)
                text = json.dumps(parsed, indent=2)
            except json.JSONDecodeError:
                pass
        lang = CONFIG_EXTENSIONS.get(ext, "")
        return f"```{lang}\n{text}\n```"
    except UnicodeDecodeError:
        return "[Binary file - cannot display as text]"


# ===== CATEGORY DETECTION =====

def get_file_category(uploaded_file) -> str:
    """Determine file category from MIME type and extension."""
    mime = uploaded_file.type or ""
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    filename_lower = uploaded_file.name.lower()

    if mime.startswith("image/") or ext in (".png",".jpg",".jpeg",".gif",".webp",".bmp"):
        return "image"
    if mime == "application/pdf" or ext == ".pdf":
        return "pdf"
    if "wordprocessingml" in mime or ext in (".docx", ".doc"):
        return "docx"
    if "spreadsheetml" in mime or "ms-excel" in mime or ext in (".xlsx",".xls"):
        return "xlsx"
    if mime in ("text/csv","text/tab-separated-values") or ext in (".csv",".tsv"):
        return "csv"
    if ext in CODE_EXTENSIONS:
        return "code"
    if ext in CONFIG_EXTENSIONS:
        return "config"
    if filename_lower in ("dockerfile","makefile","jenkinsfile","vagrantfile"):
        return "config"
    if mime.startswith("text/") or ext in TEXT_EXTENSIONS:
        return "text"
    return "unknown"


def get_file_icon(category: str) -> str:
    """Get emoji icon for file type display."""
    icons = {
        "pdf": "U0001f4d5", "docx": "U0001f4d8", "xlsx": "U0001f4ca",
        "csv": "U0001f4ca", "code": "U0001f4bb", "config": "⚙️",
        "text": "U0001f4c4", "image": "U0001f5bc️", "unknown": "U0001f4ce",
    }
    return icons.get(category, "U0001f4ce")


def get_supported_extensions() -> list:
    """Return all supported file extensions for the upload widget."""
    extensions = []
    extensions.extend(["png","jpg","jpeg","gif","webp","bmp","svg"])
    extensions.extend(["pdf","docx","doc","pptx"])
    extensions.extend(["xlsx","xls","csv","tsv"])
    extensions.extend([ext.lstrip(".") for ext in CODE_EXTENSIONS.keys()])
    extensions.extend([ext.lstrip(".") for ext in CONFIG_EXTENSIONS.keys()])
    extensions.extend([ext.lstrip(".") for ext in TEXT_EXTENSIONS.keys()])
    return sorted(set(extensions))


# ===== MAIN PROCESSING FUNCTION =====

def prepare_file_for_api(uploaded_file) -> dict:
    """
    Process any file type and return content formatted for Claude API.
    Handles: images (base64), PDFs, DOCX, XLSX, CSV, code, config, text.
    No files are stored - all processing in memory per request.
    """
    filename = uploaded_file.name
    file_size_mb = uploaded_file.size / (1024 * 1024)

    if file_size_mb > MAX_FILE_SIZE_MB:
        return {
            "type": "text",
            "text": f"[⚠️ File '{filename}' is {file_size_mb:.1f}MB - exceeds {MAX_FILE_SIZE_MB}MB limit]",
        }

    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    category = get_file_category(uploaded_file)

    # IMAGES -> Base64 for Claude Vision
    if category == "image":
        if file_size_mb > MAX_IMAGE_SIZE_MB:
            return {
                "type": "text",
                "text": f"[⚠️ Image '{filename}' is {file_size_mb:.1f}MB - exceeds {MAX_IMAGE_SIZE_MB}MB limit]",
            }
        mime_type = uploaded_file.type or "image/png"
        if mime_type == "image/jpg":
            mime_type = "image/jpeg"
        base64_data = base64.b64encode(file_bytes).decode("utf-8")
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{mime_type};base64,{base64_data}"},
        }

    # DOCUMENTS & TEXT
    if category == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif category == "docx":
        text = extract_text_from_docx(file_bytes)
    elif category == "xlsx":
        text = extract_text_from_xlsx(file_bytes)
    elif category == "csv":
        text = extract_text_from_csv(file_bytes)
    elif category == "code":
        text = extract_text_from_code(file_bytes, filename)
    elif category == "config":
        text = extract_text_from_config(file_bytes, filename)
    elif category == "text":
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = "[Binary file - cannot read as text]"
    else:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return {
                "type": "text",
                "text": f"[⚠️ Cannot read '{filename}' - unsupported binary format]",
            }

    # Truncate if too long for context window
    if len(text) > MAX_CONTEXT_CHARS:
        text = text[:MAX_CONTEXT_CHARS] + f"\n\n... [Truncated - exceeds {MAX_CONTEXT_CHARS} chars]"

    file_icon = get_file_icon(category)
    return {
        "type": "text",
        "text": f"{file_icon} **File: {filename}** ({file_size_mb:.1f}MB)\n\n{text}",
    }
